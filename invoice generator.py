import tkinter as tk
from tkinter import messagebox, filedialog
import docx
import datetime as dt
from docx2pdf import convert
import os
import customtkinter as ctk
from sqlalchemy import create_engine, Column, String, Integer, Date, Float
from sqlalchemy.orm import sessionmaker, declarative_base
from dotenv import load_dotenv

# Database setup ──────────────────────────

Base = declarative_base()

class Invoice(Base): # ORM model representing a single invoice record in the database.
    __tablename__ = "invoices" 

    id = Column(Integer, primary_key=True, autoincrement=True)
    customer_name = Column(String(64))
    customer_address = Column(String(64))
    customer_vat_id = Column(String(64))
    invoice_number = Column(String(64))
    product = Column(String(64))
    quantity = Column(Integer)
    price = Column(Float)
    vat_rate = Column(Float)
    netto = Column(Float)
    tax_amount = Column(Float)
    total = Column(Float)
    payment_terms = Column(String(64))
    invoice_date = Column(Date)

# Load DATABASE_URL from .env file to avoid hardcoding credentials in source code

load_dotenv(r'C:\Users\xxxxxxxxxxxxxxxxxxxxx\.env.txt')

# Create database engine using the connection string from the environment variable

engine = create_engine(os.getenv("my_sql_engine"))

# Create all tables defined above if they don't exist yet

Base.metadata.create_all(engine)
Session = sessionmaker(bind=engine)

class InvoiceCreator:
    """
    Builds the GUI with customtkinter, handles user input,
    fills a .docx template, converts it to PDF, and persists
    the invoice data to a SQL database via SQLAlchemy.
    """

    def __init__(self):
        self.root = ctk.CTk()
        self.root.title("Invoice PDF Creator")
        self.root.geometry('410x510')

        ctk.set_appearance_mode("dark")
        ctk.set_default_color_theme("dark-blue")

        # Input fields ───────────────
        # Each field follows the same pattern: Label -> Entry widget placed on grid

        self.Buyer_name = ctk.CTkLabel(self.root, text="Customer Name", font=('Segoe UI', 13))
        self.Buyer_name.grid(row = 0, column = 0, columnspan = 3, padx = 5)
        self.Buyer_name_entry = ctk.CTkEntry(self.root, width=400)
        self.Buyer_name_entry.grid(row = 1, column = 0, columnspan = 3, padx = 5)

        self.Buyer_address = ctk.CTkLabel(self.root, text="Customer Address", font=('Segoe UI', 13))
        self.Buyer_address.grid(row = 2, column = 0, columnspan = 3, padx = 5)
        self.Buyer_address_entry = ctk.CTkEntry(self.root, width=400)
        self.Buyer_address_entry.grid(row = 3, column = 0, columnspan = 3, padx = 5)

        self.Buyer_vat_id = ctk.CTkLabel(self.root, text="Customer VAT ID", font=('Segoe UI', 13))
        self.Buyer_vat_id.grid(row = 4, column = 0, columnspan = 3, padx = 5)
        self.Buyer_vat_id_entry = ctk.CTkEntry(self.root, width=400)
        self.Buyer_vat_id_entry.grid(row = 5, column = 0, columnspan = 3, padx = 5)

        self.Invoice_nr = ctk.CTkLabel(self.root, text="Invoice Number", font=('Segoe UI', 13))
        self.Invoice_nr.grid(row = 6, column = 0, columnspan = 3, padx = 5)
        self.Invoice_nr_entry = ctk.CTkEntry(self.root, width=400)
        self.Invoice_nr_entry.grid(row = 7, column = 0, columnspan = 3, padx = 5)

        self.Product = ctk.CTkLabel(self.root, text="Product", font=('Segoe UI', 13))
        self.Product.grid(row = 8, column = 0, columnspan = 3, padx = 5)
        self.Product_entry = ctk.CTkEntry(self.root, width=400)
        self.Product_entry.grid(row = 9, column = 0, columnspan = 3, padx = 5)

        self.values_frame = ctk.CTkFrame(self.root)
        self.values_frame.grid(row=10, column=0, columnspan=3, pady=(10, 0))

        self.Quantity = ctk.CTkLabel(self.values_frame, text="Qty", font=('Segoe UI', 13))
        self.Quantity.grid(row=0, column=0, padx=5)
        self.Quantity_entry = ctk.CTkEntry(self.values_frame, width=90)
        self.Quantity_entry.grid(row=1, column=0, padx=5)

        self.Price = ctk.CTkLabel(self.values_frame, text="Price", font=('Segoe UI', 13))
        self.Price.grid(row=0, column=1, padx=5)
        self.Price_entry = ctk.CTkEntry(self.values_frame, width=90)
        self.Price_entry.grid(row=1, column=1, padx=5)

        self.Vat_rate = ctk.CTkLabel(self.values_frame, text="Tax (%)", font=('Segoe UI', 13))
        self.Vat_rate.grid(row=0, column=2, padx=5)
        self.Vat_rate_entry = ctk.CTkEntry(self.values_frame, width=90)
        self.Vat_rate_entry.grid(row=1, column=2, padx=5)


        # Payment terms ──────────────────
        # Dict maps dropdown option -> display label + bank account number.

        self.payment_terms_dict = {
            'Immediate': {
                'Payment Terms': 'Immediate',
                'Account': 'AA123456789'
            },
            'Net7': {
                'Payment Terms': 'Net 7 days',
                'Account': 'AA123456789'
            },
            'Net30': {
                'Payment Terms': 'Net 30 days',
                'Account': 'AA123456789'
            }
        }

        self.payment_terms_label = ctk.CTkLabel(self.root, text='Payment Terms', font=('Segoe UI', 13))
        self.payment_terms_label.grid(row = 12, column = 0, columnspan = 3, padx = 5, pady = 5)

        self.payment_terms_var = tk.StringVar(self.root)
        self.payment_terms_var.set("Immediate")

        self.payment_terms_dropdown = ctk.CTkOptionMenu(
            self.root,
            variable=self.payment_terms_var,
            values=["Immediate", "Net7", "Net30"],
            width = 200,
            anchor='center',
            font = ('Segoe UI', 11))
        self.payment_terms_dropdown.grid(row = 13, column = 0, columnspan = 3, padx = 5)


        self.button_create = ctk.CTkButton(
            self.root, text='CREATE INVOICE', font=('Segoe UI', 16), width = 350,
            command=self.create_invoice
        )
        self.button_create.grid(row = 14, column = 1, padx = 15, pady = 25)

        self.root.mainloop()

    def create_invoice(self):
        """
        Main workflow triggered by the CREATE INVOICE button:
          1. Read and validate user input.
          2. Calculate net, tax, and gross amounts.
          3. Fill placeholders in the .docx template.
          4. Ask the user where to save the PDF.
          5. Convert the filled .docx to PDF, then delete the temp .docx.
          6. Persist invoice data to the database.
        """

        try:
            # Load the Word template
            doc = docx.Document(r"Invoice creator\Invoice_Dino_Template.docx")

            selected_payment_terms = self.payment_terms_dict[self.payment_terms_var.get()]

            quantity = float(self.Quantity_entry.get())
            price = float(self.Price_entry.get())
            vat_rate = float(self.Vat_rate_entry.get())

            netto = round(quantity * price,2)
            tax_amount = round(netto * (vat_rate / 100),2)
            total = netto + tax_amount

            replacements = {
                "[Customer Name]": self.Buyer_name_entry.get(),
                "[Invoice Date]": dt.datetime.today().strftime("%Y-%m-%d"),
                "[Customer Address]": self.Buyer_address_entry.get(),
                "[Customer VAT ID]": self.Buyer_vat_id_entry.get(),
                "[Invoice Number]": self.Invoice_nr_entry.get(),
                "[Vat rate]": vat_rate,
                "[Netto]": f'{netto} PLN',
                "[Tax Amount]": f'{tax_amount} PLN',
                "[Invoice Total]": f'{total} PLN',
                "[Payment Terms]": selected_payment_terms['Payment Terms'],
                "[Account]": selected_payment_terms['Account']
            }

            for paragraph in doc.paragraphs:
                for old_text, new_text in replacements.items():
                    if old_text in paragraph.text:
                        paragraph.text = paragraph.text.replace(old_text, str(new_text))

            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if "[Product]" in cell.text:
                            cell.text = cell.text.replace("[Product]", str(self.Product_entry.get()))
                        if "[Quantity]" in cell.text:
                            cell.text = cell.text.replace("[Quantity]", str(quantity))
                        if "[Price]" in cell.text:
                            cell.text = cell.text.replace("[Price]", str(price))
                        if '[Tax rate]' in cell.text:
                            cell.text = cell.text.replace('[Tax rate]', str(vat_rate))
                        if "[Total Price]" in cell.text:
                            cell.text = cell.text.replace("[Total Price]", str(total))

        except ValueError:
            messagebox.showerror('Error', message='Invalid amount or price')
            return


        invoice_number = self.Invoice_nr_entry.get()

        default_name = f'{invoice_number}.pdf'

        save_path = filedialog.asksaveasfilename(
            initialfile=default_name,
            defaultextension='.pdf',
            filetypes=[('PDF files', '*.pdf')]
        )

        if save_path:
            # docx2pdf requires an intermediate .docx file; it will be deleted after conversion
            temp_docx = save_path.replace('.pdf', '.docx')
            doc.save(temp_docx)

            try:
                convert(temp_docx, save_path)

                if os.path.exists(temp_docx):
                    os.remove(temp_docx)

                        
                data_to_save = {
                    'Customer Name': self.Buyer_name_entry.get(),
                    'Customer Address': self.Buyer_address_entry.get(),
                    'Customer VAT ID': self.Buyer_vat_id_entry.get(),
                    'Invoice Number': self.Invoice_nr_entry.get(),
                    'Product': self.Product_entry.get(),
                    'Quantity': quantity,
                    'Price': price,
                    'Vat rate': vat_rate,
                    'Netto': netto,
                    'Tax Amount': tax_amount,
                    'Invoice Total': total,
                    'Payment Terms': selected_payment_terms['Payment Terms'],
                    'Account': selected_payment_terms['Account'],
                    'Invoice Date': dt.date.today()
                }

                self.save_to_database(data_to_save)


                messagebox.showinfo("Success", f"Invoice saved as {save_path}")

                self.root.destroy()
                
            except Exception as e:
                messagebox.showerror("Error", f"Could not convert to PDF: {e}")



    def save_to_database(self, data):
        """
        Opens a new SQLAlchemy session, creates an Invoice ORM object,
        commits it, and closes the session - following the unit-of-work pattern.
        """
        session = Session()

        invoice_record = Invoice(
            customer_name=data['Customer Name'],
            customer_address=data['Customer Address'],
            customer_vat_id=data['Customer VAT ID'],
            invoice_number=data['Invoice Number'],
            product=data['Product'],
            quantity=data['Quantity'],
            price=data['Price'],
            vat_rate=data['Vat rate'],
            netto=data['Netto'],
            tax_amount=data['Tax Amount'],
            total=data['Invoice Total'],
            payment_terms=data['Payment Terms'],
            invoice_date=data['Invoice Date']
        )

        session.add(invoice_record)
        session.commit()
        session.close()


InvoiceCreator()